import os
import streamlit as st
from typing import List, Dict
import PyPDF2
import docx
import numpy as np
import pickle
from pathlib import Path

from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader, DirectoryLoader
from langchain.schema import Document
from langchain.embeddings.base import Embeddings
from langchain_community.vectorstores import FAISS
import config

class LocalEmbeddings(Embeddings):
    """Free local embeddings using sentence-transformers (CPU only)"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        # Force CPU usage to avoid CUDA compatibility issues
        import torch
        torch.set_num_threads(1)  # Prevent threading issues
        
        self.model = SentenceTransformer(model_name, device='cpu')
        # Disable torch warnings
        import warnings
        warnings.filterwarnings("ignore", category=UserWarning, module="torch")
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed search docs."""
        embeddings = self.model.encode(texts, convert_to_tensor=False)
        return embeddings.tolist()

    def embed_query(self, text: str) -> List[float]:
        """Embed query text."""
        embedding = self.model.encode([text], convert_to_tensor=False)
        return embedding[0].tolist()

class DocumentProcessor:
    def __init__(self):
        # Always use free local embeddings
        self.embeddings = LocalEmbeddings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP
        )
        
    def load_pdf(self, file_path: str) -> List[Document]:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return [Document(page_content=text, metadata={"source": file_path, "type": "pdf"})]
        except Exception as e:
            st.error(f"Error reading PDF {file_path}: {str(e)}")
            return []
    
    def load_docx(self, file_path: str) -> List[Document]:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return [Document(page_content=text, metadata={"source": file_path, "type": "docx"})]
        except Exception as e:
            st.error(f"Error reading DOCX {file_path}: {str(e)}")
            return []
    
    def load_text_file(self, file_path: str) -> List[Document]:
        """Load text from various text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                return [Document(page_content=text, metadata={"source": file_path, "type": "text"})]
        except Exception as e:
            st.error(f"Error reading text file {file_path}: {str(e)}")
            return []
    
    def process_documents(self, data_folder: str) -> List[Document]:
        """Process all documents in the data folder."""
        documents = []
        
        if not os.path.exists(data_folder):
            st.warning(f"Data folder '{data_folder}' not found. Creating it...")
            os.makedirs(data_folder, exist_ok=True)
            return documents
        
        # Walk through all files in the data folder
        for root, dirs, files in os.walk(data_folder):
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                docs = []
                
                # Process different file types
                if file_ext == '.pdf':
                    docs = self.load_pdf(file_path)
                elif file_ext == '.docx':
                    docs = self.load_docx(file_path)
                elif file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json']:
                    docs = self.load_text_file(file_path)
                
                # Add processed documents
                if docs:
                    documents.extend(docs)
                    print(f"✅ Processed {file}: {len(docs)} documents")
        
        return documents
    
    def create_vector_store(self, documents: List[Document]) -> FAISS:
        """Create and save vector store from documents using FAISS."""
        try:
            st.info("Creating FAISS vector store...")
            
            # Create FAISS vector store
            vector_store = FAISS.from_documents(
                documents=documents,
                embedding=self.embeddings
            )
            
            # Save the vector store
            vector_store.save_local(config.VECTOR_STORE_PATH)
            st.success(f"✅ Vector store saved with {len(documents)} documents")
            
        except Exception as e:
            st.error(f"Error creating vector store: {str(e)}")
            raise e
        
        return vector_store
    
    def load_existing_vector_store(self) -> FAISS:
        """Load existing FAISS vector store if it exists."""
        try:
            if os.path.exists(os.path.join(config.VECTOR_STORE_PATH, "index.faiss")):
                vector_store = FAISS.load_local(
                    config.VECTOR_STORE_PATH, 
                    self.embeddings,
                    allow_dangerous_deserialization=True
                )
                return vector_store
        except Exception as e:
            st.warning(f"Could not load existing vector store: {str(e)}")
        return None
    
    def update_documents(self) -> FAISS:
        """Update the vector store with new documents."""
        documents = self.process_documents(config.DATA_FOLDER)
        if documents:
            return self.create_vector_store(documents)
        return self.load_existing_vector_store()
